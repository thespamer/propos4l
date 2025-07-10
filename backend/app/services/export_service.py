from typing import Dict, Optional, Union, Literal
from pathlib import Path
import jinja2
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import json
import os

ExportFormat = Literal["pdf", "markdown", "html", "docx"]

class ExportService:
    def __init__(self):
        self.template_dir = Path("templates")
        self.env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(self.template_dir),
            autoescape=jinja2.select_autoescape(['html'])
        )
        
        # Load base templates
        self.html_template = self.env.get_template("proposal.html")
        self.markdown_template = self.env.get_template("proposal.md")
    
    def export(self, content: Dict, format: ExportFormat) -> Union[bytes, str]:
        """
        Export proposal to the specified format
        """
        export_methods = {
            "pdf": self._export_pdf,
            "markdown": self._export_markdown,
            "html": self._export_html,
            "docx": self._export_docx
        }
        
        if format not in export_methods:
            raise ValueError(f"Unsupported format: {format}")
        
        return export_methods[format](content)
    
    def _export_pdf(self, content: Dict) -> bytes:
        """
        Export to PDF using WeasyPrint with enhanced styling
        """
        # Generate HTML with our template
        html_content = self.html_template.render(
            content=content,
            css_class="pdf-export"
        )
        
        # Convert to PDF with WeasyPrint
        pdf = HTML(string=html_content).write_pdf()
        return pdf
    
    def _export_markdown(self, content: Dict) -> str:
        """
        Export to Markdown with proper formatting
        """
        return self.markdown_template.render(content=content)
    
    def _export_html(self, content: Dict) -> str:
        """
        Export to standalone HTML with embedded styles
        """
        return self.html_template.render(
            content=content,
            css_class="html-export",
            standalone=True
        )
    
    def _export_docx(self, content: Dict) -> bytes:
        """
        Export to Word document with proper formatting
        """
        doc = Document()
        
        # Document styles
        styles = {
            "title": {
                "font": "Arial",
                "size": 24,
                "bold": True,
                "color": RGBColor(33, 33, 33),
                "spacing": 12
            },
            "heading1": {
                "font": "Arial",
                "size": 18,
                "bold": True,
                "color": RGBColor(33, 33, 33),
                "spacing": 12
            },
            "normal": {
                "font": "Arial",
                "size": 11,
                "color": RGBColor(51, 51, 51),
                "spacing": 8
            }
        }
        
        # Add title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(content.get('title', 'Proposta Comercial'))
        self._apply_style(title_run, styles["title"])
        
        # Add metadata
        metadata = content.get('metadata', {})
        if metadata:
            meta_table = doc.add_table(rows=1, cols=2)
            meta_table.style = 'Table Grid'
            meta_cells = meta_table.rows[0].cells
            
            meta_info = [
                f"Cliente: {metadata.get('client_name', '')}",
                f"Setor: {metadata.get('industry', '')}",
                f"Data: {metadata.get('date_generated', '')}"
            ]
            meta_cells[0].text = "\n".join(meta_info)
        
        # Add sections
        sections = [
            ("Contexto", content.get('context', '')),
            ("Solução Proposta", content.get('solution', '')),
            ("Escopo do Projeto", content.get('scope', '')),
            ("Cronograma", content.get('timeline', '')),
            ("Investimento", content.get('investment', '')),
            ("Diferenciais", content.get('differentials', '')),
            ("Casos de Sucesso", content.get('cases', ''))
        ]
        
        for section_title, section_content in sections:
            if section_content:
                # Add section heading
                heading = doc.add_paragraph()
                heading_run = heading.add_run(section_title)
                self._apply_style(heading_run, styles["heading1"])
                
                # Add section content
                content_para = doc.add_paragraph()
                content_run = content_para.add_run(section_content)
                self._apply_style(content_run, styles["normal"])
        
        # Save to bytes
        from io import BytesIO
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.read()
    
    def _apply_style(self, run, style: Dict):
        """
        Apply text styling to a Word document run
        """
        font = run.font
        font.name = style["font"]
        font.size = Pt(style["size"])
        font.bold = style.get("bold", False)
        if "color" in style:
            font.color.rgb = style["color"]
